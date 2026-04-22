"""Generate Cement Fineness Prediction Report as Word document."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, highlight_rows=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)
        set_cell_shading(cell, "003366")
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
            if highlight_rows and r_idx in highlight_rows:
                set_cell_shading(cell, "D4EDDA")
            elif r_idx % 2 == 1:
                set_cell_shading(cell, "F2F6FA")
    return table


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 51, 102)


def bold_para(doc, bold_text, normal_text):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.font.bold = True
    p.add_run(normal_text)
    return p


doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ====================== TITLE PAGE ======================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Cement Mill Fineness Prediction")
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0, 51, 102)
r.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Machine Learning Analysis Report & Recommendation")
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Predictive Quality Control \u2014 Res45 (45\u03bcm Sieve Residue) Prediction")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(80, 80, 80)

for _ in range(4):
    doc.add_paragraph()

info_items = [
    ("Subject:", 'LafargeHolcim "Plants of Tomorrow" \u2014 MillMaster Digitisation'),
    ("Data Source:", "PLC/DCS sensor data \u2014 724 readings (Dec 2019 \u2013 Feb 2020)"),
    ("Date:", "April 22, 2026"),
    ("Classification:", "Internal \u2014 Confidential"),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label + " ")
    r.font.bold = True
    r.font.size = Pt(11)
    r = p.add_run(value)
    r.font.size = Pt(11)

doc.add_page_break()

# ====================== 1. EXECUTIVE SUMMARY ======================
heading(doc, "1. Executive Summary")

doc.add_paragraph(
    "We evaluated 10 machine learning models to predict cement fineness (Res45_AVG) in real-time "
    "using 8 process sensors already available from the mill's PLC/DCS system. The goal is to replace "
    "the current lab-dependent quality testing with a soft sensor that enables proactive process control."
)

doc.add_paragraph()
add_table(doc,
    ["Metric", "Value"],
    [
        ["Best Model", "LightGBM"],
        ["Accuracy", "95.88%"],
        ["Improvement Over Baseline", "+0.67% (vs 95.21% MLR)"],
        ["Models Beating Baseline", "4 out of 10"],
        ["Engineered Features", "69 (from 8 raw sensors)"],
    ])

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Recommendation: ")
r.font.bold = True
r.font.color.rgb = RGBColor(0, 128, 0)
p.add_run(
    "Deploy LightGBM model in advisory mode (Phase 1) with PLC integration via OPC-UA. "
    "Expected benefits: 5% energy reduction and 5% capacity increase per the LafargeHolcim PoC roadmap."
)

# ====================== 2. DATA OVERVIEW ======================
heading(doc, "2. Data Overview")

add_table(doc,
    ["Dataset", "Rows", "Period", "Purpose"],
    [
        ["Full Clean Data", "724", "Dec 2019 \u2013 Feb 2020", "Complete dataset"],
        ["Training Set", "514 (71%)", "Dec 2019 \u2013 Jan 2020", "Model training"],
        ["Test Set", "210 (29%)", "Feb \u2013 Aug 2020", "Model evaluation"],
    ])

doc.add_paragraph()
heading(doc, "Input Sensors (from PLC/DCS)", level=2)

add_table(doc,
    ["Sensor", "Mean", "Range", "Physical Significance"],
    [
        ["Sound Level", "0.82", "0.05 \u2013 0.91", "Mill fill level / grinding efficiency"],
        ["Production Rate", "119.88", "42.75 \u2013 155.42", "Throughput (tons/hr)"],
        ["Gypsum Feed", "5.04", "2.09 \u2013 8.35", "Additive for setting time control"],
        ["Sepol Drive Speed", "70.61", "33.94 \u2013 78.58", "Separator classifier speed"],
        ["Separator Fan", "294.33", "142.40 \u2013 312.94", "Air flow for particle separation"],
        ["Mill Feed Rate", "266.50", "46.63 \u2013 420.37", "Raw material input rate"],
        ["Bucket Elevator", "34.50", "6.03 \u2013 54.41", "Circulating load indicator"],
        ["Main Motor", "1762.34", "853.04 \u2013 1921.61", "Power draw \u2014 grinding intensity"],
    ])

doc.add_paragraph()
bold_para(doc, "Target Variable: ", "Res45_AVG (Residue on 45\u03bcm sieve, %) \u2014 Mean: 7.48%, Range: 4.4% \u2013 9.15%")
bold_para(doc, "Data Quality: ", "Zero nulls, zero duplicates. 4 mill-stop outliers removed (Sound Level < 0.1).")

# ====================== 3. FEATURE ENGINEERING ======================
doc.add_page_break()
heading(doc, "3. Feature Engineering")

doc.add_paragraph("We expanded the original 8 raw features to 69 engineered features:")

add_table(doc,
    ["Feature Type", "Count", "Purpose"],
    [
        ["Original raw sensors", "8", "Direct PLC/DCS readings"],
        ["Rolling averages (2h, 4h)", "16", "Short-term trend capture"],
        ["Rolling std deviation", "8", "Process variability/stability"],
        ["Rate of change (derivatives)", "8", "Rising/falling conditions"],
        ["Lag features (1 & 2 periods)", "16", "Process inertia effects"],
        ["Interaction terms", "7", "Physical relationships (e.g. Mill Feed x Sep Fan)"],
        ["Ratio features", "2", "Normalized efficiency metrics"],
        ["Time features (cyclical)", "4", "Shift/hour patterns"],
    ])

# ====================== 4. MODEL COMPARISON ======================
heading(doc, "4. Model Comparison Results")

add_table(doc,
    ["Rank", "Model", "Accuracy", "MAE", "RMSE", "vs Baseline", "Status"],
    [
        ["1", "LightGBM", "95.88%", "0.2904", "0.4382", "+0.67%", "BEST"],
        ["2", "XGBoost", "95.85%", "0.2885", "0.4513", "+0.64%", "BEATS BASELINE"],
        ["3", "Gradient Boosting", "95.75%", "0.2951", "0.4639", "+0.54%", "BEATS BASELINE"],
        ["4", "Stacking Ensemble", "95.52%", "0.3178", "0.5605", "+0.31%", "BEATS BASELINE"],
        ["5", "SVR (RBF)", "95.08%", "0.3505", "0.5169", "-0.13%", "CLOSE"],
        ["\u2014", "Baseline (MLR)", "95.21%", "\u2014", "\u2014", "\u2014", "ORIGINAL"],
        ["6", "Random Forest", "93.97%", "0.4226", "0.5740", "-1.24%", ""],
        ["7", "ElasticNet", "93.94%", "0.4324", "0.9340", "-1.27%", ""],
        ["8", "Lasso", "93.93%", "0.4330", "0.9183", "-1.28%", ""],
        ["9", "Ridge", "93.82%", "0.4418", "0.9907", "-1.39%", ""],
        ["10", "ANN (MLP)", "79.83%", "1.4947", "2.7654", "-15.38%", "POOR"],
    ],
    highlight_rows=[0, 1, 2, 3])

doc.add_paragraph()
bold_para(doc, "Key Finding: ",
    "Gradient boosting models (LightGBM, XGBoost, GBR) consistently outperform linear and neural network "
    "approaches. Linear models suffer from 69 highly correlated features; neural networks lack sufficient "
    "training data (510 rows).")

# ====================== 5. TOP FEATURES ======================
heading(doc, "5. Top Predictive Features")

add_table(doc,
    ["Rank", "Feature", "Importance", "Insight"],
    [
        ["1", "Sepol Drive Speed (lag 2)", "44.47%", "Separator speed from 4h ago \u2014 dominant predictor"],
        ["2", "Sepol Drive Speed (lag 1)", "8.90%", "Confirms time-lag effect on fineness"],
        ["3", "Sepol Drive Speed (change)", "4.30%", "Direction of speed change matters"],
        ["4", "Sepol Drive Speed (4h avg)", "3.28%", "Longer window average behavior"],
        ["5", "Production Rate (4h avg)", "2.65%", "Sustained throughput affects quality"],
        ["6", "Sound Level (2h avg)", "2.36%", "Mill fill level trend"],
        ["7", "Sound Level (lag 2)", "2.30%", "Historical mill fill state"],
        ["8", "Sound Level (lag 1)", "1.80%", "Recent mill fill level"],
        ["9", "Sound Level", "1.62%", "Current reading"],
        ["10", "Gypsum Feed (lag 2)", "1.38%", "Delayed additive effect"],
    ],
    highlight_rows=[0])

doc.add_paragraph()
bold_para(doc, "Physical Interpretation: ",
    "Sepol Drive Speed accounts for 61% of total predictive power across its variants. "
    "This is physically correct \u2014 the separator directly controls which particle sizes return "
    "to the mill for re-grinding. The 2-4 hour lag reflects the grinding circuit transit time.")

# ====================== 6. PLC ARCHITECTURE ======================
doc.add_page_break()
heading(doc, "6. PLC Integration \u2014 Proposed Architecture")

arch_lines = [
    "PLC/DCS (ABB MillMaster)",
    "    |  Reads 8 sensors (every 2 min)",
    "    v",
    "OPC-UA Server (data bridge)",
    "    |",
    "    v",
    "Edge Gateway / Cloud VM",
    "    |  Feature engineering (rolling avg, lags, interactions)",
    "    |  LightGBM inference (<1ms)",
    "    |  Confidence interval calculation",
    "    v",
    "Decision Engine",
    "    |  If Res45 > 8.0%  -> Recommend: Increase Sepol Drive Speed",
    "    |  If Res45 < 6.5%  -> Recommend: Decrease Separator Fan",
    "    |  If confidence low -> Flag for operator review",
    "    v",
    "OPC-UA Write-back -> PLC / HMI Dashboard",
]
for line in arch_lines:
    p = doc.add_paragraph()
    r = p.add_run(line)
    r.font.name = "Consolas"
    r.font.size = Pt(10)

# ====================== 7. ROADMAP ======================
heading(doc, "7. Implementation Roadmap")

add_table(doc,
    ["Phase", "Duration", "Activities", "Success Criteria"],
    [
        ["Phase 1: Shadow Mode", "Month 1-2",
         "Deploy model alongside existing control; Display predictions on HMI; Operators validate vs lab results",
         "Model accuracy >95% on live data"],
        ["Phase 2: Advisory Mode", "Month 3-4",
         "Model generates setpoint recommendations; Operator approves/rejects via HMI; Track acceptance rate",
         "Operator acceptance >80%; Measurable energy savings"],
        ["Phase 3: Closed-Loop", "Month 5+",
         "Auto-adjust Sepol Speed & Sep Fan within +/-5%; Operator override available; Weekly model retraining",
         "5% energy reduction; 5% capacity increase sustained"],
    ])

# ====================== 8. BUSINESS IMPACT ======================
heading(doc, "8. Expected Business Impact")

add_table(doc,
    ["Benefit", "Impact", "Source"],
    [
        ["Electrical energy reduction", "5%", "LafargeHolcim PoC"],
        ["Capacity increase", "5%", "LafargeHolcim PoC"],
        ["Lab testing reduction", "50-70% fewer manual tests", "Estimated"],
        ["Quality deviation reduction", "Eliminates off-spec production", "Real-time vs 2h lab delay"],
        ["Investment", "~80 kCHF", "LafargeHolcim roadmap"],
        ["Payback period", "4-6 months", "LafargeHolcim roadmap"],
        ["Scalability", "HIGH \u2014 all cement plants", "Uses standard PLC sensors"],
    ])

# ====================== 9. RISKS ======================
heading(doc, "9. Risks & Mitigations")

add_table(doc,
    ["Risk", "Severity", "Mitigation"],
    [
        ["Data drift (material/liner changes)", "HIGH", "Weekly retraining pipeline; drift detection"],
        ["Sepol sensor failure (44.5% dependency)", "HIGH", "Fallback model with remaining 7 sensors"],
        ["OPC-UA communication failure", "MEDIUM", "Redundant paths; auto-fallback to manual"],
        ["Operator resistance", "MEDIUM", "Phased rollout; training program"],
        ["Model overfit on historical data", "MEDIUM", "Continuous lab validation; confidence thresholds"],
        ["Safety \u2014 exceeding mechanical limits", "HIGH", "Hard PLC interlocks; capped output ranges"],
    ])

# ====================== 10. RECOMMENDATION ======================
heading(doc, "10. Recommendation")

p = doc.add_paragraph()
r = p.add_run("Proceed with LightGBM Deployment in Advisory Mode")
r.font.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()
bold_para(doc, "Why:", "")

for reason in [
    "Best accuracy at 95.88%, beating all original models",
    "Sub-millisecond inference \u2014 suitable for real-time PLC integration",
    "Feature importance aligns with physical cement grinding processes",
    "Uses only existing PLC/DCS sensors \u2014 no new hardware required",
    "Proven concept at 4 LafargeHolcim plants",
]:
    doc.add_paragraph(reason, style="List Bullet")

doc.add_paragraph()
bold_para(doc, "Immediate Next Steps:", "")

for step in [
    "IT team to set up OPC-UA data bridge between PLC/DCS and edge/cloud compute",
    "Deploy model in shadow mode for 2-month validation period",
    "Build HMI dashboard for operators to view real-time predictions",
    "Establish MLOps pipeline for automated retraining and drift detection",
]:
    doc.add_paragraph(step, style="List Number")

# ====================== SAVE ======================
out_path = r"D:\Automation\GHCP\Cement_Fineness_Prediction_Report.docx"
doc.save(out_path)
print(f"Report saved: {out_path}")
print(f"Size: {os.path.getsize(out_path) / 1024:.1f} KB")
